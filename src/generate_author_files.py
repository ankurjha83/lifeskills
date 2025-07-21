import os
from docx import Document
from docx.shared import Inches
from utils import load_json, load_markdown

GENERATED_DIR = "generated"
OUTPUT_SCRIPT = os.path.join(GENERATED_DIR, "for_author_full_script.docx")
OUTPUT_PROMPTS = os.path.join(GENERATED_DIR, "for_author_illustration_prompts.docx")

def write_script_to_docx():
    md_path = os.path.join(GENERATED_DIR, "edited_manuscript.md")
    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Edited manuscript not found at {md_path}")

    content = load_markdown(md_path)
    doc = Document()
    doc.add_heading("Final Manuscript: Life Skills Magic Series", 0)

    for line in content.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")

    doc.save(OUTPUT_SCRIPT)
    print(f"✅ Saved: {OUTPUT_SCRIPT}")

def write_prompts_to_docx():
    json_path = os.path.join(GENERATED_DIR, "illustration_prompts.json")
    if not os.path.exists(json_path):
        raise FileNotFoundError(f"Illustration prompts not found at {json_path}")

    prompts = load_json(json_path)
    doc = Document()
    doc.add_heading("Illustration Prompts for the Book", 0)

    doc.add_heading("Book-Level Illustrations", level=1)
    for key, value in prompts.items():
        if isinstance(value, str):
            doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")

    doc.add_heading("Chapter-Level Illustrations", level=1)
    for chapter in prompts.get("chapters", []):
        doc.add_paragraph(f"Chapter {chapter['chapter_number']}: {chapter['title']}", style="List Bullet")
        doc.add_paragraph(chapter["prompt"])

    doc.save(OUTPUT_PROMPTS)
    print(f"✅ Saved: {OUTPUT_PROMPTS}")

if __name__ == "__main__":
    print("📦 Generating author files...")
    write_script_to_docx()
    write_prompts_to_docx()
